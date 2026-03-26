"""Generate the Word tutorial document. Run from project root: python scripts/generate_tutorial_doc.py"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path(__file__).resolve().parent.parent / "Classroom_App_Tutorial.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    if bold:
        p.add_run(text).bold = True
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_prompt(doc, label, prompt_text):
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(prompt_text)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def build_tutorial():
    doc = Document()
    title = doc.add_heading("Tutorial: Build the Classroom Report & Analytics App with an AI Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    add_para(doc, "How to use: Open an AI coding assistant (Cursor, GitHub Copilot Chat, ChatGPT, etc.). Copy each prompt below and paste it into the chat. Apply the suggested code changes. Work through the steps in order; each step builds on the previous.")
    add_para(doc, "This guide gives you step-by-step prompts to create the app from scratch.")
    doc.add_paragraph()

    # Step 1
    add_heading(doc, "Step 1: Set up the project and config", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "Create a new folder 'Classroom App' and add: (1) pyproject.toml listing dependencies (streamlit, ollama, python-pptx, "
        "PyMuPDF, pandas, openpyxl, python-docx, plotly, fastapi, uvicorn, langgraph, langchain-core, etc.) and use uv: "
        "run `uv lock` and `uv sync`. (2) config.py with OLLAMA_HOST, OLLAMA_MODEL, tier percentages "
        "(top 20%, average 60%, low 20%), max file sizes (50 MB slides, 10 MB Excel), and TIER_LABELS dict with keys "
        "top/average/low mapped to Extension/Core/Support."
    ))
    doc.add_paragraph()

    # Step 2
    add_heading(doc, "Step 2: Create the Streamlit app skeleton", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "Create app.py: Streamlit app with set_page_config, a caption disclaimer about analytics and professional judgment, "
        "sidebar with Ollama model name input and 'Anonymize in report' checkbox. Main area: radio to switch between "
        "pages 'Upload', 'Analytics', 'Reports'. On Upload: two file uploaders (slides .pptx/.pdf, Excel .xlsx), "
        "optional text input for correct answers (comma-separated). Enforce max file sizes from config. Store uploads "
        "and settings in st.session_state. On Analytics and Reports show a short message that data will appear after "
        "parsers are built."
    ))
    doc.add_paragraph()

    # Step 3
    add_heading(doc, "Step 3: Add slide and Excel parsers", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "In parsers/: (1) slides_pptx.py: extract_text_from_pptx(path or BytesIO) returning (full_text, list of "
        "(slide_index, slide_text)); use python-pptx. Add get_poll_slides() that returns slides where first line "
        "contains 'poll' or 'question'. (2) slides_pdf.py: same for PDF using PyMuPDF page.get_text(). (3) responses.py: "
        "load_responses(file, answer_key) to read Excel with 'Student Name' column and question columns (Q1, Q2... or "
        "Q1_Selected/Q1_Correct pairs). Normalize to 0/1 correct. Support both wide format (1/0 or A/B/C with key) and "
        "Selected/Correct column pairs. normalize_responses(df) returns (df, question_columns). Add parsers/__init__.py "
        "exporting these functions."
    ))
    doc.add_paragraph()

    # Step 4
    add_heading(doc, "Step 4: Add analytics (scoring and charts)", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "In analytics/: (1) scoring.py: compute_scores(df, question_columns) = per-student score (correct/answered), "
        "assign_tiers(scores_df) using config tier percentages, get_top_n(ranked_df, n=5). All deterministic. "
        "(2) charts.py: chart_top5(top5_df, anonymize) and chart_engagement(responses_df, question_columns) using "
        "Plotly. (3) Wire app.py Upload to persist file bytes in session_state so Analytics/Reports still have files. "
        "On Analytics page: call parsers and scoring, show top-5 bar chart and engagement chart, tier counts, and "
        "ranked table. Use helpers _get_slide_file() and _get_excel_file() that fall back to persisted bytes."
    ))
    doc.add_paragraph()

    # Step 5
    add_heading(doc, "Step 5: Add Ollama client and report generation", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "In llm/ollama_client.py: check_ollama_available(host), prompt_ollama(prompt, model, host, system). "
        "OllamaClient with generate_topic_summary(lecture_text, poll_questions_text) and "
        "generate_differentiated_homework(topic_summary, tier_counts, question_specs, levels). Homework prompt must: "
        "generate only selected levels (Extension/Core/Support), use question_specs (e.g. 3 MCQ, 2 Fill in the blanks), "
        "and require MCQ answers in a separate 'Answer key' section at the end, not inline. In reports/: "
        "summary_doc.py builds Word doc from summary text; homework_doc.py parses LLM output into Extension/Core/Support "
        "sections and Answer key at the end. Coerce Excel string values to numeric in parsers and analytics to avoid "
        "str/int comparison errors."
    ))
    doc.add_paragraph()

    # Step 6
    add_heading(doc, "Step 6: Wire Reports page and homework options", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "In app.py Reports: Check Ollama availability. Ensure lecture and analytics data are loaded. Add multiselect "
        "'Generate for levels' (Extension, Core, Support) and checkboxes for question types (MCQ, Fill in the blanks, "
        "Subjective) with number inputs. On 'Generate summary' call client.generate_topic_summary and build_summary_docx; "
        "on 'Generate homework' pass topic, tier_counts, question_specs, and selected levels to "
        "generate_differentiated_homework, then build_homework_docx. Add download buttons for both .docx files. Fix "
        "Streamlit: do not assign to st.session_state keys that are used by widget key= (e.g. anonymize, ollama_model)."
    ))
    doc.add_paragraph()

    # Step 7
    add_heading(doc, "Step 7: Add guardrails and README", level=1)
    add_para(doc, "Prompt to use:")
    add_prompt(doc, "Prompt:", (
        "Add .gitignore (venv, .venv, __pycache__, .env, .DS_Store). In README: prerequisites (Python 3.10+, uv, Ollama), install "
        "steps (`uv sync`, `uv run streamlit run app.py` or API via `uv run uvicorn app:api_app`), supported Excel formats (Selected/Correct "
        "and wide), guardrails (local only, no PII in prompts, file limits, disclaimer). Optionally add templates/ with example "
        "poll_responses_example.xlsx (Student Name, Q1–Q4, 1/0) and lecture_example.pptx with a 'Poll' slide."
    ))
    doc.add_paragraph()

    # Final
    add_heading(doc, "Final check", level=1)
    add_para(doc, "Run the app: `uv run streamlit run app.py` (or the FastAPI server with `uv run uvicorn app:api_app`). "
                  "Upload the example PPT and Excel from templates/ if present, open Analytics "
                  "(confirm top-5 and engagement charts), then Reports. Generate summary and homework; download the "
                  ".docx files. Ensure Ollama is running (ollama pull llama3.2) for summary and homework generation.")
    doc.add_paragraph()

    add_para(doc, "End of tutorial.", bold=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_tutorial()
