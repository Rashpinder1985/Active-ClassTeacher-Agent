"""Word document builders for summary and homework."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build_summary_docx(summary_text: str, title: str = "Class Topic Summary", date_placeholder: bool = True) -> bytes:
    doc = Document()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_placeholder:
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    doc.add_heading("Summary", level=1)
    for para in summary_text.strip().split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(6)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_homework_docx(homework_text: str, title: str = "Differentiated Homework", date_placeholder: bool = True) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    if date_placeholder:
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    raw = (homework_text or "").strip()
    if not raw:
        doc.add_paragraph("No homework content was generated. Try generating the topic summary first, then generate homework again.")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    lines = raw.split("\n")
    level_headers = ("extension", "core", "support")
    answer_key_headers = ("answer key", "answerkey")
    current_heading = None
    current_lines: list[str] = []
    answer_key_content: list[str] = []

    def flush(heading: str, content_lines: list) -> None:
        if not heading:
            return
        block = "\n".join(content_lines).strip()
        if not block:
            return
        doc.add_heading(heading, level=1)
        for p in block.split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip())

    for line in lines:
        stripped = line.strip()
        lower = stripped.lstrip("#* ").lower()
        is_answer_key = any(lower.startswith(ak) for ak in answer_key_headers) or lower == "answer key"
        is_level = any(lower.startswith(h) for h in level_headers)

        if is_answer_key:
            if current_heading:
                flush(current_heading, current_lines)
                current_lines.clear()
            current_heading = "Answer key"
            rest = stripped.lstrip("#* ")
            for ak in ("Answer key", "Answer Key", "answer key"):
                if rest.lower().startswith(ak.lower()):
                    rest = rest[len(ak) :].lstrip(":.- ")
                    break
            if rest:
                answer_key_content.append(rest)
        elif is_level:
            if current_heading and current_heading != "Answer key":
                flush(current_heading, current_lines)
                current_lines.clear()
            elif current_heading == "Answer key":
                pass
            current_heading = (
                "Extension" if lower.startswith("extension") else "Core" if lower.startswith("core") else "Support"
            )
            if current_heading != "Answer key":
                rest = stripped.lstrip("#* ")
                for h in ("Extension", "Core", "Support"):
                    if rest.lower().startswith(h.lower()):
                        rest = rest[len(h) :].lstrip(":.- ")
                        break
                if rest:
                    current_lines.append(rest)
        else:
            if current_heading == "Answer key":
                answer_key_content.append(line)
            elif current_heading:
                current_lines.append(line)

    if current_heading and current_heading != "Answer key":
        flush(current_heading, current_lines)

    if answer_key_content:
        flush("Answer key", answer_key_content)

    if len(doc.paragraphs) <= 2:
        doc.add_heading("Homework", level=1)
        for para in raw.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
