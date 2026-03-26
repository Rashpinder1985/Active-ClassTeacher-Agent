"""Build summary report .docx from Ollama output."""
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_summary_docx(
    summary_text: str,
    title: str = "Class Topic Summary",
    date_placeholder: bool = True,
) -> bytes:
    """
    Create a Word document with title, optional date, and summary section.
    Returns document as bytes for download.
    """
    doc = Document()
    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_placeholder:
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    # Summary
    doc.add_heading("Summary", level=1)
    for para in summary_text.strip().split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(6)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
